"""Convert complex product catalogues into searchable, structured Markdown."""
from __future__ import annotations

from pathlib import Path
from dataclasses import dataclass
import re


@dataclass(frozen=True)
class CatalogFact:
    """A value stated explicitly in a parsed catalogue page."""

    entity: str
    attribute: str
    value: str
    page: int
    evidence: str


_ATTRIBUTE_RULES = (
    ("nominal_current", "Номинальный ток", "А", r"Номинальн(?:ый|ая)\s+ток(?:\s+In)?\s*(?:А|A)?\s*(.+)"),
    ("nominal_voltage", "Номинальное напряжение", "В", r"Номинальн(?:ое|ая)\s+напряжение(?:\s+Ue)?\s*(?:В|V)?\s*(.+)"),
    ("breaking_capacity", "Номинальная отключающая способность", "А", r"Номинальн(?:ая|ый)\s+отключающая способность\s*(?:А|A)?\s*(.+)"),
    ("pole_count", "Количество полюсов", "", r"Кол(?:ичество|-во)\s+полюсов\s*(?:Р|P)?\s*(.+)"),
)


def _page_sections(text: str):
    matches = list(re.finditer(r"^## Страница (\d+)\s*$", text, re.M))
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        yield int(match.group(1)), text[match.end():end]


def _entity(section: str) -> str | None:
    """Find a product-series identifier, excluding catalogue article numbers."""
    candidates = re.findall(r"\b[A-Z]{2,}\d[A-Z0-9-]*\b", section)
    candidates = [value for value in candidates if not re.fullmatch(r"A\d{5,}", value)]
    return candidates[0] if candidates else None


def extract_catalog_facts(parsed_path: Path) -> list[CatalogFact]:
    """Extract only explicit, page-level technical values from parser output.

    The extractor deliberately leaves ambiguous layouts alone. It does not infer
    missing values from product articles or neighbouring pages.
    """
    text = parsed_path.read_text(encoding="utf-8", errors="ignore")
    facts: list[CatalogFact] = []
    seen: set[tuple[str, str, str, int]] = set()
    for page, section in _page_sections(text):
        entity = _entity(section)
        if not entity:
            continue
        for line in section.splitlines():
            clean_line = " ".join(line.replace("|", " ").split())
            for attribute, _label, unit, pattern in _ATTRIBUTE_RULES:
                match = re.search(pattern, clean_line, re.I)
                if not match:
                    continue
                value = match.group(1).strip(" .;:")
                # A broken extraction can leave a single column label without a
                # value. Such rows are not facts.
                if not value or len(value) > 160 or not re.search(r"\d", value):
                    continue
                if unit and not re.search(r"(?:А|A|В|V|кА|kA)\b", value, re.I):
                    value = f"{value} {unit}"
                key = (entity, attribute, value, page)
                if key not in seen:
                    seen.add(key)
                    facts.append(CatalogFact(entity, attribute, value, page, clean_line))
                break
    return facts


_PLACEHOLDER_ANSWER_MARKERS = (
    "введите ответ",
    "нужно выделить в отдельную задачу",
    "таких таблиц у нас просто нет",
)


def _is_placeholder_answer(text: str) -> bool:
    """Not-yet-answered markers used in CNC's own tech-service Q&A template.

    Такие ответы нельзя индексировать как подтверждённый контент — цитировать
    "нужно делать с 0" пользователю как ответ нечего, это честное "ещё не
    решено", а не характеристика или правило.
    """
    lowered = text.lower()
    return any(marker in lowered for marker in _PLACEHOLDER_ANSWER_MARKERS)


@dataclass(frozen=True)
class VerifiedAnswer:
    """One curated Q&A entry from a tech-service document (Question/Technical
    Answer Word styles) — see catalog_parser.py's .docx branch below."""

    number: int
    category: str
    question: str    # первая строка Question-параграфа
    context: str      # остальные строки ("Проверено по каталогам:", "Требуется уточнить:")
    answer: str        # текст всех Technical Answer параграфов подряд, объединённый


TECH_SERVICE_ANSWER_PREFIX = "Ответ технической службы CNC (подтверждено):"


def render_verified_answers(
    entries: list[VerifiedAnswer],
    answer_prefix: str = TECH_SERVICE_ANSWER_PREFIX,
) -> list[str]:
    """Отрендерить записи «вопрос → подтверждённый ответ» в те же блоки
    "## Страница N", которые понимает knowledge_matrix._pages().

    Общая для обоих источников подтверждённых ответов: docx техслужбы CNC
    (ветка .docx ниже) и одобренная очередь ответов ИИ
    (verified_answers_queue.py) — формат должен совпадать, чтобы поиск и
    цитирование источника вели себя одинаково.

    answer_prefix обязателен к переопределению для НЕ-техслужбных источников:
    подписать одобренный инженером ответ ИИ как «ответ технической службы
    CNC» — значит приписать его авторство людям, которые его не писали, и
    поднять его в глазах читателя до уровня п.1 приоритета источников
    (ARCHITECTURE.md §5). Авторство должно оставаться честным.
    """
    parts: list[str] = []
    for entry in entries:
        parts.append(f"## Страница {entry.number}")
        parts.append(f"### {entry.category}")
        parts.append(f"Вопрос: {entry.question}")
        if entry.context:
            parts.append(entry.context)
        parts.append(f"{answer_prefix} {entry.answer}")
    return parts


def _is_verified_answers_document(document) -> bool:
    """True if this .docx uses CNC's Question/Technical Answer paragraph
    styles — checked before falling back to the plain paragraph dump so
    ordinary passports/instructions upload exactly as before."""
    style_names = {paragraph.style.name for paragraph in document.paragraphs if paragraph.style}
    return {"Question", "Technical Answer"} <= style_names


def _grouped_qa(document) -> list[tuple[str, list[str], str]]:
    """Group Heading 1 (category) / Question / Technical Answer paragraphs
    into (category, question_lines, answer_text) tuples — including
    not-yet-answered placeholder groups, unlike _extract_verified_answers().
    Shared by _extract_verified_answers() and count_short_indexed_answers()
    so the two never disagree about what counts as one Q&A group.

    "Проверено по каталогам:"/"Требуется уточнить:" are not separate
    paragraphs — they're soft line breaks inside the same Question-styled
    paragraph (paragraph.text contains embedded "\\n"), confirmed on the
    real CNC tech-service document.
    """
    groups: list[tuple[str, list[str], str]] = []
    category = ""
    question_lines: list[str] | None = None
    answer_lines: list[str] = []

    def flush() -> None:
        nonlocal question_lines, answer_lines
        if question_lines:
            answer_text = "\n".join(answer_lines).strip()
            if answer_text:
                groups.append((category, question_lines, answer_text))
        question_lines = None
        answer_lines = []

    for paragraph in document.paragraphs:
        style = paragraph.style.name if paragraph.style else ""
        text = paragraph.text.strip()
        if style == "Heading 1":
            flush()
            category = text
        elif style == "Question":
            flush()
            question_lines = [line for line in paragraph.text.split("\n") if line.strip()]
        elif style == "Technical Answer":
            if text:
                answer_lines.append(text)
        # Normal и прочие стили (преамбула, подвал "Формат ответа") — пропускаем.
    flush()
    return groups


def _extract_verified_answers(document) -> list[VerifiedAnswer]:
    """Group Heading 1 (category) / Question / Technical Answer paragraphs
    into VerifiedAnswer entries, skipping not-yet-answered placeholders."""
    entries: list[VerifiedAnswer] = []
    number = 0
    for category, question_lines, answer_text in _grouped_qa(document):
        if _is_placeholder_answer(answer_text):
            continue
        number += 1
        entries.append(VerifiedAnswer(
            number=number,
            category=category,
            question=question_lines[0].strip(),
            context="\n".join(line.strip() for line in question_lines[1:]).strip(),
            answer=answer_text,
        ))
    return entries


# Real answers in CNC's own tech-service document can legitimately be this
# short ("Да, применимо." — 14 chars) — this is a soft nudge, not a filter.
_SHORT_ANSWER_WARNING_CHARS = 40


def count_short_indexed_answers(document) -> int:
    """How many indexed (non-placeholder) answers are suspiciously short.

    _is_placeholder_answer() only recognizes the exact "not yet answered"
    phrasing CNC's template uses today — if that wording changes, such a row
    sails through and gets indexed as a confirmed fact (ARCHITECTURE.md §5:
    unverified content must never pass as one). There's no structural
    (style/formatting) signal to catch this instead — real short answers and
    placeholder rows use identical Word formatting in the source document.
    Surfaced to the uploading engineer in bot.py's reply as a nudge to check
    /documents, not a hard filter.
    """
    return sum(
        1
        for _, _, answer_text in _grouped_qa(document)
        if not _is_placeholder_answer(answer_text) and len(answer_text) < _SHORT_ANSWER_WARNING_CHARS
    )


def count_short_indexed_answers_in_file(source_path: Path) -> int:
    """count_short_indexed_answers() for callers (bot.py) that only have a
    file path — 0 for anything that isn't a recognized tech-service docx."""
    if source_path.suffix.lower() != ".docx":
        return 0
    try:
        import docx as _docx
    except ImportError:
        return 0
    document = _docx.Document(str(source_path))
    if not _is_verified_answers_document(document):
        return 0
    return count_short_indexed_answers(document)


def parse_generic_document(source_path: Path, output_path: Path, source_name: str) -> Path:
    """Convert TXT/MD/CSV/XLSX/DOCX into the same page-based Markdown format as
    parse_pdf_catalog(), so knowledge_matrix.search() finds it without any
    external vector store — retrieval stays local (ARCHITECTURE.md §2).

    Not a catalogue-grade extractor: no table-to-fact parsing beyond what
    extract_catalog_facts() already does generically on "## Страница" blocks.
    Good enough for passports, certificates and free-text notes.
    """
    suffix = source_path.suffix.lower()
    parts = [f"# {source_name}", ""]
    if suffix in {".txt", ".md"}:
        parts += ["## Страница 1", source_path.read_text(encoding="utf-8", errors="ignore")]
    elif suffix == ".csv":
        import csv as _csv
        with source_path.open(encoding="utf-8-sig", newline="") as stream:
            rows = list(_csv.reader(stream))
        parts.append("## Страница 1")
        parts.extend(" | ".join(cell.strip() for cell in row) for row in rows)
    elif suffix == ".xlsx":
        try:
            import openpyxl
        except ImportError as error:
            raise RuntimeError("Не установлен openpyxl. Выполните: python -m pip install -r requirements.txt") from error
        workbook = openpyxl.load_workbook(source_path, read_only=True, data_only=True)
        try:
            for page, sheet in enumerate(workbook.worksheets, start=1):
                parts.append(f"## Страница {page}")
                parts.append(f"### Лист: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(cell) for cell in row if cell is not None]
                    if cells:
                        parts.append(" | ".join(cells))
        finally:
            workbook.close()
    elif suffix == ".docx":
        try:
            import docx as _docx
        except ImportError as error:
            raise RuntimeError("Не установлен python-docx. Выполните: python -m pip install -r requirements.txt") from error
        document = _docx.Document(str(source_path))
        if _is_verified_answers_document(document):
            # Куратированные ответы техслужбы CNC (Question/Technical Answer
            # стили): одна "страница" на вопрос, а не весь документ единым
            # блоком — иначе knowledge_matrix.search() находит документ, но
            # snippet() показывает случайные ~24 токена, а не нужную запись.
            parts.extend(render_verified_answers(_extract_verified_answers(document)))
        else:
            parts.append("## Страница 1")
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text)
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
    else:
        raise ValueError(f"Unsupported extension for local indexing: {suffix}")
    output_path.write_text("\n".join(parts), encoding="utf-8")
    return output_path


def _cell(value: object) -> str:
    return " ".join(str(value or "").replace("\n", " ").split()).replace("|", "\\|")


def _table_to_markdown(table: list[list[object]]) -> str:
    rows = [[_cell(cell) for cell in row] for row in table if any(_cell(cell) for cell in row)]
    if len(rows) < 2 or not rows[0]:
        return ""
    width = len(rows[0])
    rows = [(row + [""] * width)[:width] for row in rows]
    return "\n".join([
        "| " + " | ".join(rows[0]) + " |",
        "| " + " | ".join("---" for _ in rows[0]) + " |",
        *["| " + " | ".join(row) + " |" for row in rows[1:]],
    ])


def parse_pdf_catalog(pdf_path: Path, output_path: Path, source_name: str) -> Path:
    """Extract page text and tables, preserving product rows for retrieval."""
    try:
        import pdfplumber
    except ImportError as error:
        raise RuntimeError("Не установлен pdfplumber. Выполните: python -m pip install -r requirements.txt") from error
    parts = [f"# Каталог CNC Electric: {source_name}", ""]
    with pdfplumber.open(pdf_path) as pdf:
        for number, page in enumerate(pdf.pages, start=1):
            parts.append(f"## Страница {number}")
            text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
            if text:
                parts.extend([text, ""])
            for index, table in enumerate(page.extract_tables(), start=1):
                markdown = _table_to_markdown(table)
                if markdown:
                    parts.extend([f"### Таблица {index}", markdown, ""])
    output_path.write_text("\n".join(parts), encoding="utf-8")
    return output_path
